from flask import Flask, request, jsonify
from flask_cors import CORS
from anonymizer import AnonymizerEngine
from entity_resolution import enhancedEntityResolutionPipeline
from vectordb import Vectordb
import re
import io
import pymupdf
import os
import csv
from langchain_core.documents import Document
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import format_document
from langchain.chains.combine_documents.base import (
    DEFAULT_DOCUMENT_PROMPT,
    DEFAULT_DOCUMENT_SEPARATOR,
)
from datetime import datetime

from litellm import completion
from routellm.controller import Controller
from datetime import datetime
from docx import Document as MSDocument

from dotenv import load_dotenv, find_dotenv
load_dotenv(find_dotenv())  # get API keys from .env file

google_api_key = os.environ.get("GOOGLE_API_KEY")
vectordb_api_key = os.environ.get("VECTOR_DB_API_KEY")
openai_api_key = os.environ["OPENAI_API_KEY"]
router_threshold = os.environ["ROUTER_THRESHOLD"]
router = os.environ["ROUTER_MODEL"]

os.system("ollama pull seeyssimon/bt4103_gguf_finance_v2")

engine = AnonymizerEngine()

vector_store = InMemoryVectorStore(
    GoogleGenerativeAIEmbeddings(
        model="models/text-embedding-004", google_api_key=google_api_key
    )
)

messages = []
file_id_tracker = {}

safety_settings = [
    {
        "category": "HARM_CATEGORY_HARASSMENT",
        "threshold": "BLOCK_NONE",
    },
    {
        "category": "HARM_CATEGORY_HATE_SPEECH",
        "threshold": "BLOCK_NONE",
    },
    {
        "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
        "threshold": "BLOCK_NONE",
    },
    {
        "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
        "threshold": "BLOCK_NONE",
    },
]

model_mapping = {
            "gemini/gemini-1.5-flash": "Gemini Model",
            "ollama_chat/seeyssimon/bt4103_gguf_finance_v2": "Finetuned Phi3.5 mini",
        }

text_splitter = RecursiveCharacterTextSplitter(chunk_size=100, chunk_overlap=20)
retriever = vector_store.as_retriever()

app = Flask(__name__)
# Allow cross-origin requests (important for frontend communication)
CORS(app, resources={r"/*": {"origins": "*"}})


def generate_output(query, origin):
    # Preprocess names
    query = enhancedEntityResolutionPipeline(query)
    # Anonymize text here
    print("Query: ", query)
    anonymized_query = engine.anonymize(query)
    print("Anonymized_query: ", anonymized_query)

    docs = retriever.invoke(anonymized_query)
    print("Length: ", len(docs))
    print("Docs list: ", docs)

    anonymized_file_text = DEFAULT_DOCUMENT_SEPARATOR.join(
        [format_document(doc, DEFAULT_DOCUMENT_PROMPT) for doc in docs]
    )
    print("Anonymized file text: ", anonymized_file_text)

    prompt = "Context: {context} \n\n Prompt: {input}".format(
        context=anonymized_file_text, input=anonymized_query
    )

    messages.append(
        {
            "role": "user",
            "content": prompt,
        }
    )

    # Use routeLLM if going through normal process or regenerating response for incorrect answer from knowledge base
    if origin is None or origin == "Vectordb":
        client = Controller(
            routers=["bert"],
            strong_model="gemini/gemini-1.5-flash",
            weak_model="ollama_chat/seeyssimon/bt4103_gguf_finance_v2",
        )

        response = client.chat.completions.create(
            model=f"router-{router}-{router_threshold}", messages=messages
        )
        model_used = response.model
        print("response model: ", response.model)

    else:  # Only use LLM if regenerating response for incorrect answer from SLM
        os.environ["GEMINI_API_KEY"] = google_api_key
        response = completion(
            model="gemini/gemini-1.5-flash",
            messages=messages,
            safety_settings=safety_settings,
        )
        # print("LLMresponse: ", response)
        model_used = "gemini/gemini-1.5-flash"

    response = response.choices[0].message.content
    model_used = model_mapping.get(model_used, model_used)

    if response == None:
        response = "Unable to generate response"

    messages.append(
        {
            "role": "assistant",
            "content": response,
        }
    )
    print("Messages: ", messages)
    # Deanonymize text here
    deanonymized_output = engine.deanonymize(response)
    print("Deanonymized_output: ", deanonymized_output)

    # Return the result as a JSON response
    return jsonify(
        {
            "original_query": query,
            "anonymized_query": anonymized_query,
            "gemini_output": response,
            "deanonymized_output": deanonymized_output,
            "model_used": model_used,
        }
    )


@app.route("/mainpipeline", methods=["POST"])
def main_pipeline():

    # Vectordb
    collection_name = "QnA"
    vectordb = Vectordb(vectordb_api_key)
    vectordb.set_threshold(0.8)
    # vectordb.create_collection(collection_name)

    # Query
    query = request.form.get("query")
    query = re.sub(r"(\d)\s+(\d)", r"\1\2", query)  # remove gaps in numbers
    print("processed query: ", query)

    # If no chat history and files uploaded, perform the knowledge base search
    if not file_id_tracker and not messages:
        # Retrieve from vectordb
        response, score = vectordb.query(query, collection_name, "answer")

        # Return response from Vectordb if similar query is found
        if response != None:
            # Convert the original query to anonymized version and store in conversation history
            query = enhancedEntityResolutionPipeline(query)
            anonymized_query = engine.anonymize(query)
            messages.append(
                {
                    "role": "user",
                    "content": anonymized_query,
                }
            )

            # Convert the original response to anonymized version and store in conversation history
            anonymized_response = engine.anonymize(response)
            messages.append(
                {
                    "role": "assistant",
                    "content": anonymized_response,
                }
            )
            print("Messages: ", messages)

            # Return the result as a JSON response
            return jsonify(
                {
                    "original_query": query,
                    "anonymized_query": anonymized_query,
                    "gemini_output": f"Similar query found in vectordb. Similarity score: {score}",
                    "deanonymized_output": response,
                    "model_used": "Vectordb",
                }
            )

    # Execute main pipeline if there is existing context or similar query not found
    response = generate_output(query, None)
    return response


@app.route("/clear_chat_history", methods=["POST"])
def clearChatHistory():
    messages.clear()
    print("Messages: ", messages)
    return jsonify({"message": "Conversation history cleared successfully."})


@app.route("/reprocessquery", methods=["POST"])
def reprocess_query():  # call the function when Vectordb output is irrelevant
    query = request.form.get("query")
    query = re.sub(r"(\d)\s+(\d)", r"\1\2", query)  # remove gaps in numbers
    print("processed query: ", query)
    # Which model we are regenerating response for
    origin = request.form.get("origin")

    # Remove the irrelevant query-answer pair from chat history
    messages.pop()
    messages.pop()

    response = generate_output(query, origin)
    return response


@app.route("/upload", methods=["POST"])
def upload_files():
    files = request.files.getlist("file")
    for file in files:
        file_name = file.filename
        file_content = file.read()
        file_extension = file_name.split(".")[-1].lower()
        print("file extension: ", file_extension)
        page_number = 0
        lst_docs = []
        file_id_tracker[file_name] = []

        if file_extension == "pdf":
            with io.BytesIO(file_content) as pdf_file:
                pdf_document = pymupdf.open(stream=pdf_file, filetype="pdf")
                anonymized_text = ""

                for page in pdf_document:
                    metadata = {}
                    metadata["source"] = file_name
                    metadata["page"] = page_number

                    text = page.get_text("text")
                    # text = re.sub(r'(\d)\s+(\d)', r'\1\2', text)  # computationally expensive, especially for large files.
                    text = enhancedEntityResolutionPipeline(text)
                    anonymized_text = engine.anonymize(text)

                    page_number += 1

                    doc = Document(metadata=metadata, page_content=anonymized_text)

                    lst_docs.append(doc)
                lst_docs = text_splitter.split_documents(lst_docs)

        elif file_extension == "csv":
            with io.StringIO(file_content.decode('utf-8')) as csv_file:
                csv_reader = csv.DictReader(csv_file)
                for i, row in enumerate(csv_reader):
                    content = "\n".join(
                            f"""{k.strip() if k is not None else k}: {v.strip()
                            if isinstance(v, str) else ','.join(map(str.strip, v))
                            if isinstance(v, list) else v}"""
                            for k, v in row.items()
                    )

                    print("CSV content: ", content)

                    doc = Document(page_content=engine.anonymize(enhancedEntityResolutionPipeline(content)), metadata={"source": file_name, "row": i})
                    print("CSV document: ", doc)
                    lst_docs.append(doc)

        elif file_extension == "docx":
            with io.BytesIO(file_content) as docx_file:
                ms_document = MSDocument(docx_file)
                print("ms doc: ", ms_document)
                full_text = []
                for paragraph in ms_document.paragraphs:
                    full_text.append(paragraph.text)
                print("full text: ", full_text)
            full_doc = "\n".join(full_text)
            print("full_doc: ", full_doc)
            doc = Document(page_content=engine.anonymize(enhancedEntityResolutionPipeline(full_doc)) , metadata={"source": file_name})
            lst_docs = [doc]
            lst_docs = text_splitter.split_documents(lst_docs)
                    
        lst_ids = vector_store.add_documents(lst_docs)
        file_id_tracker[file_name].extend(lst_ids)

        print(f"Processed and anonymized: {file_name}")
    # print(anonymized_file_text)
    return jsonify(
        {"message": f"{len(files)} file(s) uploaded and processed successfully."}
    )


@app.route("/clear_anonymized_text", methods=["POST"])
def clear_anonymized_text():
    file_name = request.json.get("fileName")
    if file_name in file_id_tracker:
        print("File name: ", file_name)
        print("File id tracker: ", file_id_tracker)
        print("File name present in file id tracker: ", file_name in file_id_tracker)

        # Delete the file content from the in-memory vector db
        vector_store.delete(file_id_tracker[file_name])
        print("Vector store: ", vector_store)
        # Delete the file info from the file_id_tracker
        del file_id_tracker[file_name]
        print("File id tracker: ", file_id_tracker)

        retriever = vector_store.as_retriever()

        # print(anonymized_file_text)
        return jsonify(
            {"message": f"Anonymized text for {file_name} cleared successfully."}
        )
    else:
        return jsonify({"message": "File not found."})


@app.route("/handle_feedback", methods=["POST"])
def handle_feedback():
    query = request.json.get("query")
    answer = request.json.get("answer")
    document = [
        {
            "query": query,
            "answer": answer,
            "date": datetime.now()
        }
    ]
    collection_name = "QnA"
    vectordb = Vectordb(vectordb_api_key)
    print("Answer", answer)
    vectordb.upload_docs(collection_name, document, "query")
    return jsonify({"message": "Uploaded query-answer to vectorDB."})

@app.route("/get_recent_queries", methods=["POST"])
def get_recent_queries():
    vectordb = Vectordb(vectordb_api_key)
    recent_queries = vectordb.get_recent_queries(4)
    return jsonify(
      {
      "query_1": recent_queries[0],
      "query_2" : recent_queries[1],
      "query_3" : recent_queries[2],
      "query_4" : recent_queries[3],
     }
    )


if __name__ == "__main__":
    app.run(debug=True, use_reloader=False, host = '0.0.0.0', port='5000')
